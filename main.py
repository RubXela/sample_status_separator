import os
import sys
import tkinter
from tkinter import PhotoImage, filedialog, messagebox, ttk

import docx
import pandas as pd


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath('.')
    return os.path.join(base_path, relative_path)

custom_img = resource_path('blan.png')


class FileEditorApp(tkinter.Tk):
    def __init__(self):
        super().__init__()
        self.title('Распределение шаблонов по статусам')
        img = PhotoImage(file=custom_img)
        self.tk.call('wm', 'iconphoto', self._w, img)
        self.geometry('1100x650')
        self.configure(highlightbackground='gray', highlightthickness=1)
        self.resizable(True, True)
        self.file_workplace = tkinter.LabelFrame(
            self, width=250, height=150, text='Работа с файлами',)
        self.file_workplace.place(x=10, y=5)
        self.upload_button = tkinter.Button(
            self, width=30, height=1, text='Загрузить СЗ',
            highlightbackground='gray', highlightthickness=2, border=1,
            command=lambda: get_file_to_upload())
        self.upload_button.place(x=20, y=30)
        self.save_button = tkinter.Button(
            self, width=30, height=1, text='Сохранить файл',
            highlightbackground='gray', highlightthickness=2, border=1,
            command=lambda: save_to_file())
        self.save_button.place(x=20, y=60)
        self.label_start_sell = tkinter.Label(
            self, text='')
        self.label_start_sell.place(x=20, y=100)
        self.label_end_sell = tkinter.Label(
            self, text='')
        self.label_end_sell.place(x=20, y=125)

        self.doc_workplace = tkinter.LabelFrame(
            self, width=300, height=150, text='Информация о СЗ')
        self.doc_workplace.place(x=270, y=5)
        self.label_all_status = tkinter.Label(
            self, text='')
        self.label_all_status.place(x=280, y=25)
        self.label_potential_status = tkinter.Label(
            self, text='')
        self.label_potential_status.place(x=280, y=50)
        self.label_renew_status = tkinter.Label(
            self, text='')
        self.label_renew_status.place(x=280, y=75)
        self.label_lost_status = tkinter.Label(
            self, text='')
        self.label_lost_status.place(x=280, y=100)
        self.label_undefined_status = tkinter.Label(
            self, text='')
        self.label_undefined_status.place(x=280, y=125)

        self.headers = ''
        self.potential_member = []
        self.lost_member = []
        self.renew_member = []
        self.undefined_member = []

        def save_to_file():
            if self.headers != '':
                saved_filename = filedialog.asksaveasfilename(
                initialdir='C:\\Users\\user\\Desktop',
                title='Сохранить файл',
                filetype=[('Excel', '*.xlsx')])
                try:
                    if saved_filename != '':
                        info_data = pd.DataFrame.from_dict({
                            'Информация о СЗ': [
                                self.label_all_status['text'],
                                self.label_potential_status['text'],
                                self.label_renew_status['text'],
                                self.label_lost_status['text'],
                                self.label_start_sell['text'],
                                self.label_end_sell['text'],
                            ]
                        })
                        potential = get_formated_data(
                            self.potential_member, self.headers)
                        renew = get_formated_data(
                            self.renew_member, self.headers)
                        lost = get_formated_data(
                            self.lost_member, self.headers)
                        if not '.xlsx' in saved_filename:
                            saved_filename = saved_filename + '.xlsx'
                        with pd.ExcelWriter(saved_filename) as writer:
                            if potential is not None:
                                potential.to_excel(
                                    writer,
                                    sheet_name='ПЧК',
                                    index=False)
                            if renew is not None:
                                renew.to_excel(
                                    writer,
                                    sheet_name='РП, МП, ПП',
                                    index=False)
                            if lost is not None:
                                lost.to_excel(
                                    writer,
                                    sheet_name='БЧК',
                                    index=False)
                            info_data.to_excel(
                                writer,
                                sheet_name='Информация о СЗ',
                                index=False
                            )
                except FileNotFoundError:
                    return
                except Exception:
                    messagebox.showerror(
                        'Информация', 'Выбран неверный формат файла')

        def get_formated_data(table_rows: list, table_headers: list):
            if table_rows != []:
                data = pd.DataFrame(table_rows)
                data.columns = table_headers
                return data
            else:
                return None

        def get_file_to_upload():
            filename = filedialog.askopenfilename(
                initialdir='C:\\Users\\{oper}\\Desktop',
                title='Выберите файл',
                filetypes=[('Word', '*.docx')])
            create_table_data(filename)

        def create_table_data(name: str):
            try:
                if name != '':
                    clear_table()
                    self.potential_member.clear()
                    self.renew_member.clear()
                    self.lost_member.clear()
                    self.headers = ''
                    file = docx.Document(name)
                    for table in file.tables:
                        df = [[
                            '' for i in range(len(table.columns))
                             ] for j in range(len(table.rows))]
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                if cell.text:
                                    df[i][j] = cell.text
                    dataframe = pd.DataFrame(df[:][1:], index=None, columns=df[1])
                    list_data = dataframe.to_numpy().tolist()
                    self.headers = list_data[0]
                    self.table['column'] = self.headers
                    self.table['show'] = 'headings'
                    for index, column in enumerate(self.table['columns']):
                        width = len(list_data[1][index])
                        self.table.heading(column, anchor=tkinter.CENTER,
                                           text=column)
                        anchor = tkinter.CENTER
                        if width > 20:
                            width = 200
                            anchor = 'w'
                        self.table.column(index, minwidth=15, anchor=anchor,
                                          width=width)
                    potential_counter = 0
                    lost_counter = 0
                    renew_counter = 0
                    undefined_counter = 0
                    parent_potential = ''
                    parent_lost = ''
                    parent_renew = ''
                    parent_undefined = ''
                    status_col = ''
                    for index, row in enumerate(list_data):
                        if index == 0:
                            for desc in row:
                                if desc.strip().lower() == 'статус':
                                    status_col = desc
                        status = (
                            str(dataframe[status_col][index]).strip().lower()
                            ).strip(',')
                        if '\n' in row[0]:
                            row[0] = (row[0].split('\n')[0]
                            + row[0].split('\n')[1])
                        if 'пчк' in status:
                            if parent_potential == '':
                                parent_potential = self.table.insert(
                                    '', 'end', value='ПЧК')
                            self.potential_member.append(row[:])
                            self.table.insert(
                                parent_potential, 'end', value=row)
                            potential_counter += 1
                        elif 'бчк' in status:
                            if parent_lost == '':
                                parent_lost = self.table.insert(
                                    '', 'end', value='БЧК')
                            self.lost_member.append(row[:])
                            self.table.insert(parent_lost, 'end', value=row)
                            lost_counter += 1
                        elif (not 'бчк' in status
                            and not 'пчк' in status
                            and not 'рп, мп, пп' in status
                            and not 'рп,мп,пп' in status
                            and not 'рп' in status
                            and not 'мп' in status
                            and not 'пп' in status
                            and index != 0):
                            if parent_undefined == '':
                                parent_undefined = self.table.insert(
                                    '', 'end', value='БЕЗ-СТАТУСА')
                            self.undefined_member.append(row[:])
                            self.table.insert(parent_undefined, 'end', value=row)
                            undefined_counter += 1
                        elif ('рп, мп, пп' in status
                            or 'рп,мп,пп' in status
                            or 'рп' in status
                            or 'мп' in status
                            or 'пп' in status
                            or status in ('рп', 'мп', 'пп')
                            ):
                            if parent_renew == '':
                                parent_renew = self.table.insert(
                                    '', 'end', value='РП-МП-ПП')
                            self.renew_member.append(row[:])
                            self.table.insert(parent_renew, 'end', value=row)
                            renew_counter += 1
                    all_count = (potential_counter + undefined_counter +
                                 lost_counter + renew_counter)
                    fname = name.split('/')[-1]
                    self.title(f'Распределение шаблонов по статусам: {fname}')
                    self.label_all_status['text'] = (
                        f'Общее количество шаблонов: {all_count}')
                    self.label_potential_status['text'] = (
                        f'ПЧК: {potential_counter}')
                    self.label_renew_status['text'] = (
                        f'РП, МП, ПП: {renew_counter}')
                    self.label_lost_status['text'] = f'БЧК: {lost_counter}'
                    self.label_undefined_status['text'] = (
                        f'БЕЗ СТАТУСА: {undefined_counter}')
                    for index, line in enumerate(file.paragraphs):
                        line = str(line.text).strip().lower()
                        if 'дата введения услуги' in line:
                            if '\n' in line:
                                line = line.split('\n')[1]
                            if line.split(':')[1] != '':
                                date = line.split(':')[1]
                                self.label_start_sell['text'] = (
                                    f'Начало продажи: {date}')
                            else:
                                date = file.paragraphs[index + 1].text
                                self.label_start_sell['text'] = (
                                    f'Начало продажи: {date}')
                        elif 'дата окончания продажи' in line:
                            if '\n' in line:
                                line = line.split('\n')[1]
                            if line.split(':')[1] != '':
                                date = line.split(':')[1]
                                self.label_end_sell['text'] = (
                                    f'Окончание продажи: {date}')
                            else:
                                date = file.paragraphs[index + 1].text
                                self.label_end_sell['text'] = (
                                    f'Окончание продажи: {date}')                      
            except FileNotFoundError:
                return
            except Exception as e:
                messagebox.showerror('Ошибка', f'Описание: {e}')

        self.table_workplace = tkinter.LabelFrame(
            self, text='Таблица служебной записки')
        self.table_workplace.place(x=10, y=160, relwidth=.985, relheight=.7)
        self.table = CustomTreeView(self.table_workplace, show='headings')
        self.table.place(x=10, y=10, relwidth=.985, relheight=.95)
        treescrolly = tkinter.Scrollbar(
            self.table, orient='vertical', command=self.table.yview)
        self.table.configure(yscrollcommand=treescrolly.set)
        treescrolly.pack(side="right", fill="y")

        def clear_table():
            self.table.delete(*self.table.get_children())
            self.table['show'] = ''
            return None
        

class CustomTreeView(ttk.Treeview):
    def __init__(self, master, **kw):
        super().__init__(master, **kw)
        self.bind('<space>', self.on_select)

    def on_select(self, event):

        selected_reg = self.identify_region(event.x, event.y)
        if selected_reg != 'cell':
            return
        column = self.identify_column(event.x)
        column_idx = int(column[1:]) - 1
        selected_row_id = self.focus()
        selected_values = self.item(selected_row_id)
        column_box = self.bbox(selected_row_id, column)
        if (selected_values.get('values')[0] == 'ПЧК'
            or selected_values.get('values')[0] == 'РП-МП-ПП'
            or selected_values.get('values')[0] == 'БЧК'):
            return
        self.clipboard_clear()
        select_text = selected_values.get('values')[column_idx]
        entry_edit = ttk.Entry(self, width=len(str(select_text)))
        entry_edit.editing_column_index = column_idx
        entry_edit.editing_item_row_id = selected_row_id
        entry_edit.insert(0, select_text)
        entry_edit.select_range(0, 'end')
        self.clipboard_append(select_text)
        entry_edit.focus()
        entry_edit.bind('<FocusOut>', self.unfocus)
        entry_edit.place(x=column_box[0], y=column_box[1], w=column_box[2],
                         h=column_box[3])
        
    def unfocus(self, event):
        event.widget.destroy()


if __name__ == "__main__":
    app = FileEditorApp()
    app = ttk.Entry(app)
    app.mainloop()
